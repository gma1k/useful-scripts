#!/usr/bin/env python3

import fitz
from docx import Document
from docx.shared import Inches
import io
from PIL import Image
import os  # <-- Add this

pdf_path = input("Enter the path to the PDF file: ").strip()
docx_path = input("Enter the path for the output Word file (.docx): ").strip()

doc = Document()

pdf = fitz.open(pdf_path)

for page_number in range(len(pdf)):
    page = pdf[page_number]

    text = page.get_text("text")
    if text:
        for line in text.split("\n"):
            doc.add_paragraph(line)

    for img_index, img in enumerate(page.get_images(full=True)):
        xref = img[0]
        base_image = pdf.extract_image(xref)
        image_bytes = base_image["image"]
        image_ext = base_image["ext"]
        image = Image.open(io.BytesIO(image_bytes))

        tmp_path = f"temp_image_{page_number}_{img_index}.{image_ext}"
        image.save(tmp_path)
        doc.add_picture(tmp_path, width=Inches(5))
        image.close()
        os.remove(tmp_path)  # <-- works now

doc.save(docx_path)
print(f"PDF successfully converted to {docx_path}")
